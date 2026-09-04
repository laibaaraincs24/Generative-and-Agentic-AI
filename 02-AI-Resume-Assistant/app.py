import io
import json
import os
import re
from typing import Any, Dict, List

import streamlit as st
from docx import Document
from google import genai
from google.genai import types
from pypdf import PdfReader


# ============================================================
# CONFIG
# ============================================================

DEFAULT_MODEL = "gemini-3.7-flash"
MAX_RESUME_CHARS = 30000


# ============================================================
# PAGE
# ============================================================

st.set_page_config(
    page_title="ResumeLens - ATS Resume Analyzer",
    page_icon="📄",
    layout="wide",
)

st.title("📄 ResumeLens")
st.caption(
    "Upload your resume to get an ATS-readiness score, keyword analysis, "
    "and AI-powered improvement suggestions."
)


# ============================================================
# RESUME TEXT EXTRACTION
# ============================================================

def extract_pdf_text(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages)


def extract_docx_text(file_bytes: bytes) -> str:
    document = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

    # Include text from tables because some resumes put contact/skills
    # information inside tables.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    return "\n".join(paragraphs)


def extract_txt_text(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="replace")


def extract_resume_text(uploaded_file) -> str:
    file_bytes = uploaded_file.getvalue()
    suffix = uploaded_file.name.lower().rsplit(".", 1)[-1]

    if suffix == "pdf":
        return extract_pdf_text(file_bytes)
    if suffix == "docx":
        return extract_docx_text(file_bytes)
    if suffix == "txt":
        return extract_txt_text(file_bytes)

    raise ValueError("Unsupported file type. Please upload PDF, DOCX, or TXT.")


# ============================================================
# DETERMINISTIC ATS CHECKS
# ============================================================

SECTION_PATTERNS = {
    "Contact information": [
        r"\bemail\b", r"@", r"\bphone\b", r"\bmobile\b", r"\blinkedin\b"
    ],
    "Summary / Objective": [
        r"\bsummary\b", r"\bprofessional summary\b", r"\bobjective\b",
        r"\bprofile\b"
    ],
    "Experience": [
        r"\bexperience\b", r"\bwork experience\b", r"\bemployment\b",
        r"\bprofessional experience\b"
    ],
    "Education": [
        r"\beducation\b", r"\bacademic\b", r"\bdegree\b", r"\bbachelor\b",
        r"\bmaster\b", r"\buniversity\b"
    ],
    "Skills": [
        r"\bskills\b", r"\btechnical skills\b", r"\bcore competencies\b",
        r"\btechnologies\b"
    ],
}


def count_words(text: str) -> int:
    return len(re.findall(r"\b[\w+#.-]+\b", text))


def extract_keywords(text: str) -> List[str]:
    """Return likely technical/professional keywords for simple ATS analysis."""
    stopwords = {
        "the", "and", "for", "with", "from", "that", "this", "your", "you",
        "are", "was", "were", "have", "has", "had", "will", "our", "their",
        "into", "using", "used", "about", "over", "under", "after", "before",
        "through", "then", "than", "also", "such", "work", "working",
        "experience", "skills", "education", "resume", "cv", "role", "job",
        "team", "teams", "project", "projects", "responsible", "responsibilities",
        "including", "include", "strong", "good", "excellent", "professional",
    }

    tokens = re.findall(r"\b[A-Za-z][A-Za-z0-9+#.-]{2,}\b", text.lower())
    frequency = {}
    for token in tokens:
        if token not in stopwords and not token.isdigit():
            frequency[token] = frequency.get(token, 0) + 1

    return [
        word for word, _ in sorted(
            frequency.items(), key=lambda item: (-item[1], item[0])
        )[:30]
    ]


def calculate_basic_ats_score(resume_text: str, job_description: str = "") -> Dict[str, Any]:
    text = resume_text.lower()
    words = count_words(resume_text)

    section_hits = {}
    for section, patterns in SECTION_PATTERNS.items():
        section_hits[section] = any(re.search(pattern, text) for pattern in patterns)

    section_score = sum(section_hits.values()) / len(section_hits) * 30

    contact_score = 0
    if re.search(r"[\w.+-]+@[\w-]+\.[\w.-]+", resume_text):
        contact_score += 5
    if re.search(r"(?:\+?\d[\d\s().-]{7,}\d)", resume_text):
        contact_score += 5
    if "linkedin.com" in text:
        contact_score += 5

    length_score = 15
    if words < 150:
        length_score = 5
    elif words < 300:
        length_score = 10
    elif words > 1400:
        length_score = 8

    formatting_score = 20
    formatting_flags = []

    if "|" in resume_text:
        formatting_score -= 3
        formatting_flags.append("Table/column-like separators detected; keep layouts simple.")
    if re.search(r"[^\x00-\x7F]", resume_text):
        formatting_score -= 1
    if text.count("•") > 0 or text.count("- ") > 2:
        formatting_score += 0
    else:
        formatting_flags.append("Use clear bullet points for achievements and responsibilities.")

    keyword_score = 0
    matched_keywords = []
    missing_keywords = []

    if job_description.strip():
        jd_keywords = extract_keywords(job_description)
        resume_words = set(re.findall(r"\b[A-Za-z][A-Za-z0-9+#.-]{2,}\b", text))
        matched_keywords = [k for k in jd_keywords if k in resume_words]
        missing_keywords = [k for k in jd_keywords if k not in resume_words]
        keyword_score = (
            min(1.0, len(matched_keywords) / max(1, len(jd_keywords))) * 30
        )
        # Replace the generic keyword portion with job-specific matching.
        raw_score = section_score + contact_score + length_score + formatting_score
        final_score = round(min(100, raw_score * 0.7 + keyword_score), 1)
    else:
        # Generic ATS readiness: measurable structure + contact + length + formatting.
        final_score = round(
            min(100, section_score + contact_score + length_score + formatting_score),
            1,
        )

    return {
        "score": final_score,
        "word_count": words,
        "sections": section_hits,
        "formatting_flags": formatting_flags,
        "matched_keywords": matched_keywords,
        "missing_keywords": missing_keywords,
    }


# ============================================================
# GEMINI ANALYSIS
# ============================================================

def get_api_key() -> str:
    # Streamlit Cloud: st.secrets["GEMINI_API_KEY"]
    # Local: environment variable GEMINI_API_KEY
    try:
        secret_key = st.secrets.get("GEMINI_API_KEY", "")
    except Exception:
        secret_key = ""

    return secret_key or os.getenv("GEMINI_API_KEY", "")


def analyze_with_gemini(
    resume_text: str,
    job_description: str,
    basic_analysis: Dict[str, Any],
) -> Dict[str, Any]:
    api_key = get_api_key()
    if not api_key:
        raise RuntimeError(
            "Gemini API key is missing. Add GEMINI_API_KEY to Streamlit Secrets "
            "or set it as an environment variable."
        )

    client = genai.Client(api_key=api_key)

    prompt = f"""
You are an expert ATS resume reviewer and technical recruiter.

Analyze the resume below. Your score must be an ATS-READINESS ESTIMATE, not a
claim about any particular company's proprietary ATS algorithm.

Return ONLY valid JSON with exactly these top-level keys:
- ats_score: integer from 0 to 100
- verdict: one of ["Excellent", "Good", "Needs Improvement", "Weak"]
- summary: short string
- strengths: array of 3 to 6 concise strings
- improvements: array of 5 to 10 concise strings
- missing_keywords: array of strings
- keyword_matches: array of strings
- formatting_issues: array of strings
- impact_issues: array of strings
- recommended_sections: array of strings

Scoring guidance:
- 20% contact and basic information
- 20% standard ATS-friendly section structure
- 20% skills and keyword alignment
- 20% measurable achievements / strong action verbs
- 10% readability and concise wording
- 10% formatting / ATS parseability

If a job description is provided, prioritize alignment with that job.
Do NOT invent experience, qualifications, employers, dates, certifications,
or skills. Recommendations must be truthful and based on the supplied text.

A deterministic pre-check produced this information:
{json.dumps(basic_analysis, indent=2)}

JOB DESCRIPTION:
{job_description.strip() or "No job description provided. Perform a general ATS-readiness review."}

RESUME:
{resume_text[:MAX_RESUME_CHARS]}
"""

    response = client.models.generate_content(
        model=DEFAULT_MODEL,
        contents=prompt,
        config=types.GenerateContentConfig(
            response_mime_type="application/json",
            temperature=0.2,
            max_output_tokens=4000,
        ),
    )

    if not response.text:
        raise RuntimeError("Gemini returned an empty response.")

    try:
        result = json.loads(response.text)
    except json.JSONDecodeError as exc:
        raise RuntimeError("Gemini returned invalid JSON. Please try again.") from exc

    return result


# ============================================================
# UI
# ============================================================

with st.sidebar:
    st.header("Analysis options")
    st.write(
        "For the most useful keyword score, paste the job description for the "
        "role you are targeting."
    )
    job_description = st.text_area(
        "Job description (optional)",
        height=260,
        placeholder="Paste the target job description here...",
    )

    st.divider()
    st.info(
        "Privacy note: the resume text is sent to Gemini for analysis. "
        "Do not upload sensitive documents you are not comfortable sending "
        "to the configured AI service."
    )

uploaded_file = st.file_uploader(
    "Upload your resume",
    type=["pdf", "docx", "txt"],
    help="Supported formats: PDF, DOCX, TXT",
)

analyze_button = st.button(
    "🔍 Analyze Resume",
    type="primary",
    use_container_width=True,
)

if analyze_button:
    if uploaded_file is None:
        st.warning("Please upload a resume first.")
        st.stop()

    try:
        resume_text = extract_resume_text(uploaded_file)
    except Exception as exc:
        st.error(f"Could not read the resume: {exc}")
        st.stop()

    if not resume_text.strip():
        st.error(
            "No readable text was found. If this is a scanned/image-only PDF, "
            "please upload a text-based PDF or DOCX."
        )
        st.stop()

    if len(resume_text) > MAX_RESUME_CHARS:
        st.warning(
            f"The extracted resume is longer than {MAX_RESUME_CHARS:,} characters. "
            "Only the first portion will be sent to Gemini."
        )

    basic = calculate_basic_ats_score(resume_text, job_description)

    with st.spinner("Analyzing your resume with Gemini Flash..."):
        try:
            ai_result = analyze_with_gemini(
                resume_text,
                job_description,
                basic,
            )
        except Exception as exc:
            st.error(f"Analysis failed: {exc}")
            st.stop()

    ai_score = int(max(0, min(100, int(ai_result.get("ats_score", basic["score"])))))

    if ai_score >= 85:
        score_label = "Excellent"
    elif ai_score >= 70:
        score_label = "Good"
    elif ai_score >= 50:
        score_label = "Needs Improvement"
    else:
        score_label = "Weak"

    st.divider()
    st.subheader("ATS Readiness")

    col1, col2, col3 = st.columns(3)
    col1.metric("ATS Score", f"{ai_score}/100")
    col2.metric("Resume Words", f'{basic["word_count"]:,}')
    col3.metric("Verdict", ai_result.get("verdict", score_label))

    st.progress(ai_score / 100)

    st.caption(
        "This is an ATS-readiness estimate based on resume structure, content, "
        "keyword alignment, impact, and formatting. It is not an official score "
        "from a specific employer's ATS."
    )

    tab1, tab2, tab3, tab4 = st.tabs(
        ["✨ Improvements", "🔑 Keywords", "📐 ATS Checks", "📝 Summary"]
    )

    with tab1:
        st.markdown("### What to improve")
        for item in ai_result.get("improvements", []):
            st.markdown(f"- {item}")

        st.markdown("### Impact / achievement issues")
        impact_issues = ai_result.get("impact_issues", [])
        if impact_issues:
            for item in impact_issues:
                st.markdown(f"- {item}")
        else:
            st.success("No major impact-writing issues were identified.")

        st.markdown("### Formatting issues")
        formatting_issues = ai_result.get("formatting_issues", [])
        if formatting_issues:
            for item in formatting_issues:
                st.markdown(f"- {item}")
        else:
            st.success("No major formatting issues were identified.")

    with tab2:
        if job_description.strip():
            st.markdown("### Matched keywords")
            matched = ai_result.get("keyword_matches", []) or basic["matched_keywords"]
            if matched:
                st.write(", ".join(matched))
            else:
                st.info("No strong keyword matches were identified.")

            st.markdown("### Keywords to consider adding")
            missing = ai_result.get("missing_keywords", []) or basic["missing_keywords"]
            if missing:
                st.write(", ".join(missing))
            else:
                st.success("No obvious missing keywords were identified.")
        else:
            st.info(
                "Add a job description in the sidebar to get job-specific "
                "keyword matching."
            )

    with tab3:
        st.markdown("### Standard section checks")
        for section, present in basic["sections"].items():
            if present:
                st.success(f"✓ {section}")
            else:
                st.warning(f"⚠ {section}")

        st.markdown("### Basic formatting checks")
        if basic["formatting_flags"]:
            for flag in basic["formatting_flags"]:
                st.warning(flag)
        else:
            st.success("No basic formatting flags detected.")

    with tab4:
        st.markdown("### AI summary")
        st.write(ai_result.get("summary", "No summary returned."))

        st.markdown("### Strengths")
        for item in ai_result.get("strengths", []):
            st.markdown(f"- {item}")

        recommended = ai_result.get("recommended_sections", [])
        if recommended:
            st.markdown("### Recommended sections")
            for item in recommended:
                st.markdown(f"- {item}")

    with st.expander("View extracted resume text"):
        st.text(resume_text[:MAX_RESUME_CHARS])
